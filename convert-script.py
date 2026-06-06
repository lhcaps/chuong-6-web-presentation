"""
Convert SCRIPT_THUYET_TRINH.md to Word document with beautiful formatting.
Run: python convert-script.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page margins ───────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)    # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.9)
section.bottom_margin  = Inches(0.9)

# ─── Colour palette ──────────────────────────────────────────────────────────
DARK   = RGBColor(0x11, 0x18, 0x27)   # #111827 — tiêu đề, nội dung chính
MUTED  = RGBColor(0x37, 0x41, 0x51)   # #374151 — nội dung phụ, gợi ý
ACCENT = RGBColor(0x25, 0x63, 0xEB)   # #2563EB — số slide, tiêu đề phụ
TEAL   = RGBColor(0x0F, 0x76, 0x6E)   # #0F766E — phần/chương
RED    = RGBColor(0xB4, 0x53, 0x09)   # #B45309 — cảnh báo
GRAY   = RGBColor(0x9C, 0xA1, 0xAA)   # cho đường kẻ

# ─── Helper: set paragraph shading ──────────────────────────────────────────
def shade_paragraph(para, hex_fill: str):
    """Add background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)

# ─── Helper: add a horizontal rule ──────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ─── Helper: heading styles ──────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = DARK
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size  = Pt(15)
    run.font.bold  = True
    run.font.color.rgb = ACCENT
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = TEAL
    return p

# ─── Helper: body text ──────────────────────────────────────────────────────
def body(doc, text, italic=False, muted=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.bold  = bold
    run.font.color.rgb = MUTED if muted else DARK
    return p

def talk(doc, text):
    """Script line — normal weight."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.15)
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.color.rgb = DARK
    return p

def action(doc, text, color=None):
    """[TƯƠNG TÁC] / [GỢI Ý] / [CHUYỂN CẢNH] line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold  = True
    run.font.color.rgb = color or ACCENT
    return p

def note(doc, text):
    """Timing / slide number note."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = GRAY
    return p

# ─── Helper: slide header box ───────────────────────────────────────────────
def slide_header(doc, number, title, section_name):
    """Shaded box for slide title."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Inches(0)

    shade_paragraph(p, "F0F4FF")  # very light blue bg

    run_num = p.add_run(f"  TRANG {number}  ")
    run_num.font.size  = Pt(9)
    run_num.font.bold  = True
    run_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # bg chip for number — use shading on a sub-run via separate para

    run_title = p.add_run(f"{title.upper()}")
    run_title.font.size  = Pt(13)
    run_title.font.bold  = True
    run_title.font.color.rgb = DARK

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Inches(0)
    shade_paragraph(p2, "F0F4FF")
    run_sec = p2.add_run(f"  {section_name}  ")
    run_sec.font.size  = Pt(9)
    run_sec.font.color.rgb = ACCENT
    run_sec.font.bold  = True

# ─────────────────────────────────────────────────────────────────────────────
#  COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(80)
p_title.paragraph_format.space_after  = Pt(6)
r = p_title.add_run("SCRIPT THUYẾT TRÌNH")
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = DARK

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(4)
r = p_sub.add_run("CHƯƠNG 6: TRIỂN KHAI VIỆC TẠO LẬP DOANH NGHIỆP")
r.font.size  = Pt(16)
r.font.color.rgb = ACCENT

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_info.paragraph_format.space_before = Pt(4)
p_info.paragraph_format.space_after  = Pt(4)
r = p_info.add_run("Nhóm 3 — Giảng viên: Đỗ Đình Thanh — Môn: Khởi nghiệp")
r.font.size  = Pt(11)
r.font.color.rgb = MUTED

p_time = doc.add_paragraph()
p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_time.paragraph_format.space_before = Pt(4)
p_time.paragraph_format.space_after  = Pt(30)
r = p_time.add_run("50 trang  ·  ~18-20 phút")
r.font.size  = Pt(10)
r.font.italic = True
r.font.color.rgb = GRAY

add_hr(doc)

# ─── Navigation guide ────────────────────────────────────────────────────────
h2(doc, "BẢNG HƯỚNG DẪN SỬ DỤNG")
add_hr(doc)

body(doc,
     "Mỗi trang ghi rõ: nội dung nói, gợi ý tương tác, thời gian, chuyển cảnh.\n"
     "Các nhãn trong script có ý nghĩa:")
body(doc, "[TƯƠNG TÁC]  →  Chỗ dừng lại hỏi lớp hoặc cho thảo luận. Có thể bỏ qua nếu không đủ thời gian.")
body(doc, "[GỢI Ý]       →  Mẹo nói hoặc cách nhấn mạnh một ý.")
body(doc, "[CHUYỂN CẢNH]  →  Câu nói chuyển sang slide kế tiếp một cách mượt mà.")
body(doc, "⏱             →  Thời gian ước tính cho slide đó.")
body(doc, "")
body(doc,
     "Với phong cách hướng nội: nói chậm, ít di chuyển, tập trung vào nội dung rõ ràng. "
     "Không cần oai lối. Không cần cố tỏ ra năng động. Cứ là chính mình, nói những gì mình hiểu.",
     muted=True)

add_hr(doc)

# ─── PHẦN 1 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 1  ·  MỞ ĐẦU  (Slide 1–5)  ·  ~2 phút")
add_hr(doc)

# ── Slide 1 ──
slide_header(doc, "01", "Cover", "Opening")
note(doc, "⏱  Khoảng 16 giây")
body(doc, "Tiêu đề: Chương 6: Triển khai việc tạo lập doanh nghiệp")
body(doc, "")
body(doc, "Nói (giọng bình tĩnh, chậm rãi):")
talk(doc, '"Xin chào thầy và các bạn. Em là Lê Huy, nhóm trưởng nhóm 3. '
       'Hôm nay nhóm em sẽ trình bày Chương 6: Triển khai việc tạo lập doanh nghiệp, môn Khởi nghiệp."')
body(doc, "")
body(doc, "Nói tiếp (giọng có trọng lượng hơn):")
talk(doc, '"Câu hỏi đặt ra là: có business plan rồi thì sáng mai làm gì? '
       'Câu trả lời nằm trong Chương 6 — biến kế hoạch thành hệ thống vận hành có thể kiểm soát."')
action(doc, "[CHUYỂN CẢNH] →  \"Trước tiên, cho em giới thiệu 7 thành viên trong nhóm.\"")
note(doc, "Tip hướng nội: đứng vững, nhìn thẳng, không vội. Im lặng 2 giây sau câu hỏi cũng là cách tạo ấn tượng tốt.")

# ── Slide 2 ──
slide_header(doc, "02", "Thành viên nhóm", "Opening")
note(doc, "⏱  Khoảng 16 giây")
body(doc, "Nói (chỉ điểm nhanh từng người, không đọc chi tiết):")
talk(doc, '"Đây là 7 thành viên của nhóm 3. Em là Lê Huy, nhóm trưởng, phụ trách điều phối flow và Q&A cuối cùng. '
       'Các thành viên còn lại: Tuấn Hưng, Huỳnh Hải, Nguyễn Khánh, Nguyễn Hiền, Long Gia Hảo và Huỳnh Sơn Hải. '
       'Mỗi người phụ trách một phần nội dung riêng trong pipeline này."')
body(doc, "")
talk(doc, '"Nhưng quan trọng là cả nhóm đều nắm flow tổng, nên thầy có thể hỏi bất kỳ thành viên nào trong Q&A."')
action(doc, "[CHUYỂN CẢNH] →  \"Vào bài chính thôi — vì sao Chương 6 lại quan trọng?\"")
note(doc, "Tip hướng nội: giới thiệu thành viên bằng tên thật, không cần đọc mã số. Lướt nhanh, đừng dừng lại quá lâu.")

# ── Slide 3 ──
slide_header(doc, "03", "Vì sao Chương 6 quan trọng?", "Opening")
note(doc, "⏱  Khoảng 20 giây")
body(doc, "Nói (giọng rõ ràng, nhấn từng ý):")
talk(doc, '"Doanh nghiệp thường không chết vì thiếu ý tưởng. Mà chết vì triển khai thiếu hệ thống."')
body(doc, "")
talk(doc, '"Cụ thể, có ba lớp việc mà bất kỳ doanh nghiệp nào cũng phải làm — không thể né."')
body(doc, "")
body(doc, "Bấm reveal từng ý, nói rõ từng ý:")
talk(doc, '"Thứ nhất: Pháp lý. Không đăng ký đúng, gặp rủi ro ngay từ đầu."')
talk(doc, '"Thứ hai: Vận hành. Không phân vai rõ, founder ôm hết mà việc vẫn không xong."')
talk(doc, '"Thứ ba: Số liệu. Không đọc được dòng tiền, có doanh thu cũng có thể phá sản."')
action(doc, "[CHUYỂN CẢNH] →  \"Ba lớp việc đó nằm trong toàn bộ chương. Đây là bản đồ tổng.\"")
note(doc, "Tip hướng nội: câu đầu là thesis statement — nói chậm, rõ từng từ, có ngừng nhịp. Không cần phải lớn tiếng hay hoang mang.")

# ── Slide 4 ──
slide_header(doc, "04", "Bản đồ toàn chương", "Opening")
note(doc, "⏱  Khoảng 20 giây")
body(doc, "Nói (đi từ trái sang phải, mỗi phần 2-3 giây):")
talk(doc, '"Chương 6 không phải một danh sách dài. Nó là một hệ thống gồm 5 phần liên kết với nhau."')
body(doc, "")
talk(doc, '"Pháp lý — chọn loại hình, đăng ký, chi phí, rủi ro. Đây là kiến trúc nền tảng."')
talk(doc, '"Vận hành — vai trò founder, nhân sự lõi, quy trình tối thiểu, SOP. Đây là bản đồ quyền lực."')
talk(doc, '"Thị trường — MVP, kiểm chứng khách hàng, kênh bán, phễu chuyển đổi. Đây là nơi học từ thực tế."')
talk(doc, '"Tài chính — chi phí khởi tạo, dòng tiền, hòa vốn, phân bổ nguồn lực. Đây là radar."')
talk(doc, '"Kiểm soát — rủi ro và KPI. Đây là hệ thống cảnh báo sớm."')
action(doc, "[CHUYỂN CẢNH] →  \"Và cả hệ thống này bắt đầu từ: từ ý tưởng sang doanh nghiệp thật.\"")
note(doc, "Tip hướng nội: đọc từng phần như đang kể chuyện, không phải đọc báo. Giọng đều, không quá nhanh.")

# ── Slide 5 ──
slide_header(doc, "05", "Từ ý tưởng sang doanh nghiệp thật", "Opening")
note(doc, "⏱  Khoảng 17 giây")
body(doc, "Nói:")
talk(doc, '"Ý tưởng chỉ có giá trị khi chuyển thành ba thứ: quyết định, quy trình và chỉ số."')
body(doc, "")
talk(doc, '"Quyết định — chọn cái gì làm trước, cái gì sau. Không chọn là không có gì."')
talk(doc, '"Quy trình — làm thế nào để chất lượng lặp lại được, không phụ thuộc trí nhớ."')
talk(doc, '"Chỉ số — biết đúng sai ở đâu, để điều chỉnh kịp thời."')
body(doc, "")
talk(doc, '"Cả chương hôm nay trả lời ba câu hỏi đó."')
action(doc, "[CHUYỂN CẢNH] →  \"Và trước khi làm bất cứ điều gì, cần một checklist để không bỏ sót.\"")
note(doc, "Tip hướng nội: câu cuối \"cả chương trả lời ba câu hỏi đó\" là câu hay quên — nhắc lại nó giúp lớp thấy bài có kết nối.")

# ─── PHẦN 2 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 2  ·  PHÁP LÝ  (Slide 6–14)  ·  ~4 phút")
add_hr(doc)

# ── Slide 6 ──
slide_header(doc, "06", "Checklist triển khai tổng quát", "Planning")
note(doc, "⏱  Khoảng 19 giây")
body(doc, "Nói (giọng nhẹ nhàng, như chia sẻ kinh nghiệm):")
talk(doc, '"Slide này là bản đồ để cả nhóm không bị lạc. Năm mục không thể bỏ sót:"')
body(doc, "")
talk(doc, '"1. Loại hình pháp lý — chọn trước khi làm bất cứ điều gì khác."')
talk(doc, '"2. Hồ sơ đăng ký — giấy tờ pháp lý phải xong trước khi nhận tiền."')
talk(doc, '"3. MVP — sản phẩm đủ để học từ khách thật, không cần hoàn hảo."')
talk(doc, '"4. Kênh bán — biết bán cho ai và bằng cách nào."')
talk(doc, '"5. Dòng tiền — biết đang tiêu bao nhiêu và còn bao lâu."')
body(doc, "")
talk(doc, '"Mỗi mục không cần hoàn hảo ngay, nhưng phải có người chịu trách nhiệm và deadline."')
action(doc, "[CHUYỂN CẢNH] →  \"Bắt đầu từ mục đầu tiên: Chọn mô hình kinh doanh.\"")

# ── Slide 7 ──
slide_header(doc, "07", "Chọn mô hình kinh doanh", "Planning")
note(doc, "⏱  Khoảng 19 giây")
body(doc, "Nói (đơn giản, không thuật ngữ):")
talk(doc, '"Mô hình kinh doanh nghe to, nhưng thực ra chỉ cần trả lời ba câu hỏi:"')
body(doc, "")
talk(doc, '"Câu 1: Ai trả tiền? — Khách hàng mục tiêu là ai?"')
talk(doc, '"Câu 2: Trả vì giá trị nào? — Họ nhận được gì mà sẵn sàng trả tiền?"')
talk(doc, '"Câu 3: Chi phí chính nằm ở đâu? — Tiền vào bằng cơ chế nào, tiền ra bằng cơ chế nào?"')
action(doc, "[CHUYỂN CẢNH] →  \"Ba câu hỏi đó giúp xác định mô hình. Nhưng trước khi bán, phải xác định loại hình pháp lý.\"")
note(doc, "Tip hướng nội: đừng dùng thuật ngữ như 'revenue model'. Nói tiếng Việt đơn giản, ai nghe cũng hiểu.")

# ── Slide 8 ──
slide_header(doc, "08", "Chọn loại hình pháp lý", "Legal")
note(doc, "⏱  Khoảng 22 giây")
body(doc, "Nói (cân bằng, không thiên vị loại nào):")
talk(doc, '"Ba loại hình phổ biến ở Việt Nam: Hộ kinh doanh, Công ty TNHH và Công ty cổ phần."')
body(doc, "")
talk(doc, '"Hộ kinh doanh — phù hợp quy mô nhỏ, một chủ, ít rủi ro, bán lẻ nhỏ. '
       'Trách nhiệm vô hạn — tài sản cá nhân có thể bị ảnh hưởng."')
talk(doc, '"Công ty TNHH — phù hợp khi có từ 2 người trở lên, muốn tách biệt tài sản cá nhân và doanh nghiệp. '
       'Quản trị vừa phải, dễ gọi vốn hơn."')
talk(doc, '"Công ty cổ phần — phù hợp khi có tham vọng gọi vốn lớn, muốn mời cổ đông. '
       'Quản trị chặt hơn."')
action(doc, "[TƯƠNG TÁC TÙY CHỌN] →  \"Nếu một nhóm 3 bạn muốn bán đồ ăn online quanh trường, nên chọn loại hình nào?\"")
action(doc, "[CHUYỂN CẢNH] →  \"Sau khi chọn loại hình, cần chuẩn bị hồ sơ.\"")

# ── Slide 9 ──
slide_header(doc, "09", "Decision tree: hộ kinh doanh hay doanh nghiệp", "Legal")
note(doc, "⏱  Khoảng 27 giây")
body(doc, "Nói (từ tốn, rõ ràng từng bước):")
talk(doc, '"Đây là cây quyết định để chọn đúng loại hình. Đi theo từng nhánh:"')
body(doc, "")
talk(doc, '"Bước 1: Quy mô và trách nhiệm — bạn muốn tách biệt tài sản cá nhân không? Có thì đi tiếp."')
talk(doc, '"Bước 2: Một chủ hay nhiều founder? Một chủ, bán nhỏ, ít rủi ro → Hộ kinh doanh."')
talk(doc, '"Bước 3: Nhiều founder hoặc có vốn góp? → Công ty TNHH."')
talk(doc, '"Bước 4: Có rủi ro trách nhiệm cao hoặc cần gọi vốn lớn? → Công ty cổ phần."')
action(doc, "[CHUYỂN CẢNH] →  \"Sau khi chọn loại hình, cần chuẩn bị hồ sơ.\"")

# ── Slide 10 ──
slide_header(doc, "10", "Hồ sơ đăng ký", "Legal")
note(doc, "⏱  Khoảng 19 giây")
body(doc, "Nói (tập trung vào những chỗ hay sai):")
talk(doc, '"Hồ sơ đăng ký là bản cam kết đầu tiên về danh tính, vốn, ngành nghề và người đại diện. '
       'Năm phần cần chuẩn bị:"')
body(doc, "")
talk(doc, '"Tên doanh nghiệp — phải độc đáo, không trùng, không vi phạm nhãn hiệu."')
talk(doc, '"Địa chỉ trụ sở — phải là địa chỉ thật, có thể nhận thư tín."')
talk(doc, '"Ngành nghề — khai đúng và đủ. Ngành nghề có điều kiện phải đáp ứng tiêu chuẩn trước khi đăng ký."')
talk(doc, '"Vốn điều lệ — khai thật. Không khai cao hơn thực tế góp."')
talk(doc, '"Người đại diện — phải là người thực sự điều hành, chịu trách nhiệm pháp lý."')
action(doc, "[CHUYỂN CẢNH] →  \"Hồ sơ xong rồi, quy trình đăng ký như thế nào?\"")

# ── Slide 11 ──
slide_header(doc, "11", "Quy trình đăng ký", "Legal")
note(doc, "⏱  Khoảng 23 giây")
body(doc, "Nói (nhấn mạnh bước 5):")
talk(doc, '"Đăng ký không kết thúc ở giấy phép. Đây là năm bước:"')
body(doc, "")
talk(doc, '"1. Chọn tên — tra cứu trước trên Cổng thông tin quốc gia."')
talk(doc, '"2. Chuẩn bị hồ sơ — theo checklist slide trước."')
talk(doc, '"3. Nộp đăng ký — nộp online qua Cổng Dịch vụ công, mất 3-5 ngày làm việc."')
talk(doc, '"4. Nhận kết quả — Giấy chứng nhận đăng ký doanh nghiệp."')
talk(doc, '"5. Hậu đăng ký — ĐÂY LÀ PHẦN NHIỀU BẠN BỎ QUA. Sau khi có giấy phép, '
       'vẫn còn: đăng ký thuế, mua hóa đơn, mở tài khoản ngân hàng doanh nghiệp, '
       'đăng ký chữ ký số và quy trình lưu chứng từ."')
action(doc, "[CHUYỂN CẢNH] →  \"Tất cả đều có chi phí. Chi phí pháp lý ban đầu như thế nào?\"")

# ── Slide 12 ──
slide_header(doc, "12", "Chi phí pháp lý ban đầu", "Legal")
note(doc, "⏱  Khoảng 19 giây")
body(doc, "Nói (giọng chia sẻ):")
talk(doc, '"Chi phí pháp lý nhỏ hơn chi phí sửa sai sau khi vận hành rất nhiều."')
body(doc, "")
talk(doc, '"Thiết bị chiếm phần lớn nhất — vì thường cần máy tính, phần mềm, công cụ."')
talk(doc, '"Mặt bằng — nếu cần không gian vật lý."')
talk(doc, '"Dự phòng — luôn để ra 15-20% cho chi phí phát sinh không lường trước."')
talk(doc, '"Marketing và pháp lý — chiếm phần nhỏ nhưng không được bỏ qua."')
body(doc, "")
body(doc, "Ý chính: đưa pháp lý vào ngân sách khởi tạo, không coi là việc phụ.", muted=True)
action(doc, "[CHUYỂN CẢNH] →  \"Và nếu không làm đúng, rủi ro pháp lý nào thường gặp nhất?\"")

# ── Slide 13 ──
slide_header(doc, "13", "Rủi ro pháp lý thường gặp", "Legal")
note(doc, "⏱  Khoảng 22 giây")
body(doc, "Nói (giọng nhẹ nhưng thẳng thắn):")
talk(doc, '"Rủi ro pháp lý thường đến từ chi tiết nhỏ bị xem nhẹ. Bốn lỗi phổ biến:"')
body(doc, "")
talk(doc, '"1. Tên gây nhầm lẫn — trùng hoặc tương tự tên đã đăng ký → bị yêu cầu đổi, mất thời gian."')
talk(doc, '"2. Ngành nghề có điều kiện — không kiểm tra trước → đăng ký xong rồi mới biết không đủ điều kiện."')
talk(doc, '"3. Hợp đồng miệng — nhận tiền nhưng không có văn bản → không có căn cứ pháp lý khi tranh chấp."')
talk(doc, '"4. Không lưu chứng từ — chi tiêu không có hóa đơn → không khấu trừ được thuế, không chứng minh chi phí."')
action(doc, "[CHUYỂN CẢNH] →  \"Từ pháp lý, chuyển sang phần tiếp theo: Vận hành.\"")

# ── Slide 14 ──
slide_header(doc, "14", "Case mini: chọn sai loại hình", "Legal")
note(doc, "⏱  Khoảng 27 giây")
body(doc, "Nói (kể chuyện, có ngữ điệu — đây là story, không phải bài giảng):")
talk(doc, '"Đây là một case giả lập nhưng rất thực tế. Nhiều bạn trong lớp mình có thể gặp."')
body(doc, "")
talk(doc, '"Ban đầu: Một nhóm 3 bạn bán đồ ăn online. Chọn Hộ kinh doanh vì nhanh và đơn giản."')
talk(doc, '"Sau đó: Có người quen muốn góp vốn 30 triệu để mở rộng. Nhưng Hộ kinh doanh chỉ có một chủ — '
       'không thể ghi nhận cổ phần."')
talk(doc, '"Vấn đề: Không rõ quyền, trách nhiệm và phần lợi nhuận của người góp vốn. '
       'Ai chịu lỗi? Ai được chia lời? Thỏa thuận miệng không có giá trị pháp lý."')
talk(doc, '"Kết quả: Hoặc đồng ý mơ hồ rồi xung đột sau, hoặc phải chuyển đổi loại hình — tốn chi phí và thời gian."')
action(doc, "[CHUYỂN CẢNH] →  \"Từ pháp lý, chuyển sang: Vận hành.\"")

# ─── PHẦN 3 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 3  ·  VẬN HÀNH  (Slide 15–20)  ·  ~3 phút")
add_hr(doc)

slides_vh = [
    (15, "Xây dựng bộ máy vận hành", "Operations", 22,
     'Bộ máy vận hành là cách doanh nghiệp biến lời hứa thành kết quả lặp lại.\n\n'
     '"Câu hỏi cốt lõi: Ai làm gì, ai kiểm tra, ai chịu trách nhiệm?"\n\n'
     '"Một doanh nghiệp mới không cần bộ máy lớn, nhưng cần phân vai rõ. '
     'Không thì founder ôm hết việc, làm đến kiệt sức mà việc vẫn không xong."'),
    (16, "Vai trò founder", "Operations", 20,
     '"Founder không chỉ nghĩ ý tưởng. Founder thiết kế hệ thống ra quyết định."\n\n'
     'Bốn vai trò lớn nhất của founder giai đoạn đầu:\n'
     '1. Giữ tầm nhìn — biết đang đi đâu, không bị cuốn theo công việc hàng ngày.\n'
     '2. Chốt ưu tiên — quyết định làm cái gì trước, cái gì sau, cái gì không làm.\n'
     '3. Bảo vệ dòng tiền — không chi tiêu hoang phí khi chưa có doanh thu ổn định.\n'
     '4. Gỡ nút thắt — khi team bị stuck, founder là người mở khóa.\n\n'
     '"Founder giai đoạn đầu phải làm nhiều việc, nhưng không được nhầm bận rộn với vận hành tốt."'),
    (17, "Vai trò nhân sự lõi", "Operations", 20,
     '"Đội lõi phải phủ đủ bốn năng lực: làm ra, bán được, giao được và đo được."\n\n'
     'Product — làm ra sản phẩm/dịch vụ.\n'
     'Sales — bán được cho khách hàng.\n'
     'Operations — giao được sản phẩm/dịch vụ đúng chất lượng.\n'
     'Finance — đo được bằng số, kiểm soát được tiền.\n\n'
     '[TƯƠNG TÁC] Hỏi lớp: nếu thiếu một trong bốn năng lực thì sẽ ra sao? '
     'Thiếu Sales: làm ra nhưng không bán được. '
     'Thiếu Ops: bán được nhưng giao không nổi. '
     'Thiếu Finance: có doanh thu nhưng mất kiểm soát tiền.'),
    (18, "Sơ đồ tổ chức giai đoạn đầu", "Operations", 22,
     '"Tổ chức nhỏ vẫn cần đường báo cáo rõ. '
     'Mục tiêu không phải giống công ty lớn, mà là giảm nhầm lẫn khi có vấn đề."\n\n'
     'Founder ở trung tâm. Bốn trụ cột: Product, Sales, Ops, Finance báo cáo lên founder.\n\n'
     'Nguyên tắc: mỗi việc phải có một owner rõ ràng. '
     'Không ai làm việc của ai, nhưng ai cũng biết việc đó thuộc về ai.'),
    (19, "Quy trình vận hành tối thiểu", "Operations", 23,
     '"Quy trình tối thiểu giúp chất lượng không phụ thuộc vào trí nhớ."\n\n'
     'Lấy ví dụ đơn giản: một đơn hàng đi qua năm bước:\n'
     'Nhận đơn → Xử lý → Giao hàng → Thu tiền → Ghi nhận dữ liệu.\n\n'
     'Mất một bước? → Mất dữ liệu hoặc mất tiền.\n\n'
     'Ví dụ: nhận đơn, chuẩn bị, giao hàng — nhưng quên ghi nhận. '
     'Tháng sau không biết mình đã bán bao nhiêu, thu được bao nhiêu.'),
    (20, "Standard Operating Procedure", "Operations", 23,
     '"SOP tốt là đủ ngắn để dùng thật, đủ rõ để người mới làm được."\n\n'
     'Một SOP giai đoạn đầu chỉ cần năm phần:\n'
     '1. Mục đích — tại sao cần quy trình này?\n'
     '2. Input — đầu vào là gì?\n'
     '3. Các bước — liệt kê theo thứ tự, mỗi bước một dòng.\n'
     '4. Owner — ai chịu trách nhiệm?\n'
     '5. Output — kết quả đầu ra là gì, bằng chứng hoàn thành là gì?\n\n'
     '"Đừng biến SOP thành văn bản dài 20 trang. Một trang là đủ giai đoạn đầu."'),
]

for num, title, section_name, seconds, content in slides_vh:
    slide_header(doc, str(num).zfill(2), title, section_name)
    note(doc, f"⏱  Khoảng {seconds} giây")
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("[TƯƠNG TÁC]"):
            action(doc, line, ACCENT)
        elif line.startswith("[CHUYỂN CẢNH]"):
            action(doc, line, TEAL)
        elif line.startswith('"'):
            talk(doc, line)
        else:
            body(doc, line)

# ─── PHẦN 4 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 4  ·  THỊ TRƯỜNG  (Slide 21–28)  ·  ~3.5 phút")
add_hr(doc)

slides_tt = [
    (21, "Sản phẩm/dịch vụ tối thiểu", "Market", 22,
     '"Sản phẩm tối thiểu không phải bản cẩu thả. '
     'Mà là bản đủ để học từ khách thật."\n\n'
     'MVP cần đáp ứng ba điều:\n'
     '1. Một nhu cầu rõ — khách hàng THẬT SỰ cần, không phải mình nghĩ họ cần.\n'
     '2. Một lời hứa rõ — nói được mình giải quyết vấn đề gì, bằng cách nào.\n'
     '3. Một cách đo phản hồi — khách phản hồi bằng hành vi, không chỉ bằng lời nói.\n\n'
     '[TƯƠNG TÁC] Hỏi lớp: nếu các bạn làm MVP bán cà phê, MVP đó là gì? '
     'Không phải một quán đầy đủ. Có thể chỉ là một cart cà phê handmade với một vài món.'),
    (22, "Lean startup và MVP", "Market", 23,
     '"MVP giúp giảm rủi ro bằng cách học sớm trước khi đốt nhiều tiền."\n\n'
     'Chu kỳ Lean Startup gồm bốn bước lặp:\n'
     '1. Giả thuyết — mình nghĩ khách hàng sẽ làm X.\n'
     '2. MVP — tạo bản thử nhỏ nhất để kiểm chứng giả thuyết đó.\n'
     '3. Khách dùng — đưa MVP đến tay khách thật, quan sát hành vi.\n'
     '4. Dữ liệu — thu thập dữ liệu, rồi quyết định: tiếp tục hay thay đổi.\n\n'
     '"Lean không phải làm ít cho nhanh. Mà là học nhanh với chi phí hợp lý."'),
    (23, "Build Measure Learn", "Market", 25,
     '"Chu kỳ học nhanh biến triển khai thành quá trình kiểm chứng liên tục."\n\n'
     'Build — tạo bản thử nhỏ nhất có thể.\n'
     'Measure — đo hành vi khách, không chỉ hỏi ý kiến.\n'
     'Learn — rút ra kết luận, thay đổi quyết định nếu cần.\n'
     'Decide — quyết định: tiếp tục theo hướng này hay pivot?\n\n'
     '[TƯƠNG TÁC] Câu hỏi: nếu khách khen sản phẩm nhưng không mua, đó là tín hiệu gì? '
     'Đó là tín hiệu PHẢI ĐO HÀNH VI, không chỉ nghe lời khen. '
     'Khách nói thích nhưng hành vi cho thấy không đủ nhu cầu thật.'),
    (24, "Kiểm chứng khách hàng", "Market", 23,
     '"Khách hàng thật được kiểm chứng bằng hành vi, không chỉ bằng khảo sát."\n\n'
     'Thang kiểm chứng từ thấp đến cao:\n'
     '1. Nói thích — khảo sát, hỏi ý kiến. Dễ nhưng dễ sai.\n'
     '2. Đăng ký — khách để lại thông tin. Tốt hơn nói thích.\n'
     '3. Dùng thử — khách dùng thử miễn phí hoặc trả tiền nhỏ. Tín hiệu mạnh.\n'
     '4. Trả tiền thật — mua lần đầu. Tín hiệu rất mạnh.\n'
     '5. Mua lại — quay lại mua tiếp. Tín hiệu mạnh nhất.\n\n'
     '"Càng xuống dưới, tín hiệu càng đáng tin. Đừng dừng ở bước 1 hoặc 2."'),
    (25, "Kênh bán hàng ban đầu", "Market", 22,
     '"Kênh đầu tiên nên là nơi HỌC được nhiều nhất, '
     'không nhất thiết là nơi MỞ RỘNG lớn nhất."\n\n'
     'Bốn tiêu chí chọn kênh ban đầu:\n'
     '1. Nhanh phản hồi — biết khách nghĩ gì trong vài ngày, không phải vài tuần.\n'
     '2. Chi phí thấp — giai đoạn đầu không đủ ngân sách cho quảng cáo lớn.\n'
     '3. Dễ đo — đo được số người tiếp cận, số người mua.\n'
     '4. Có thể lặp lại — quy trình bán được nhắc lại mà không cần nhiều thay đổi.\n\n'
     'Ví dụ: bán trực tiếp cho bạn bè/trường học, mạng xã hội, cộng đồng nhỏ. '
     'Không phải chạy quảng cáo Facebook ngay.'),
    (26, "Marketing launch plan", "Market", 23,
     '"Launch plan không phải làm ồn. Mà là tạo một đợt thử có mục tiêu đo được."\n\n'
     'Bốn yếu tố cần chuẩn bị:\n'
     '1. Một thông điệp — nói một điều, nói rõ, nhớ được.\n'
     '2. Một tệp khách — biết mình bán cho ai, không phải ai cũng.\n'
     '3. Một ưu đãi — khuyến khích hành động đầu tiên.\n'
     '4. Một dashboard đo — trước khi launch, biết mình sẽ đo cái gì.\n\n'
     'Đề xuất launch nhỏ trong 7 đến 14 ngày. Có mục tiêu cụ thể.'),
    (27, "Phễu khách hàng", "Market", 22,
     '"Phễu cho biết khách rơi ở đâu và cần sửa bước nào."\n\n'
     'Năm bước trong phễu: Tiếp cận → Quan tâm → Dùng thử → Mua lần đầu → Mua lại.\n\n'
     'Dựa vào biểu đồ, bước nào rơi nhiều nhất?\n'
     'Ví dụ: 100 người tiếp cận nhưng chỉ 10 người mua. '
     'Rơi ở bước quan tâm → vấn đề nằm ở thông điệp hoặc giá. '
     'Rơi ở bước dùng thử → vấn đề nằm ở trải nghiệm.\n\n'
     '"Không cần phễu phức tạp. Chỉ cần nhìn tỷ lệ rơi qua từng bước."'),
    (28, "Biểu đồ chuyển đổi", "Market", 22,
     '"Một biểu đồ tốt giúp cả nhóm nhìn cùng một sự thật."\n\n'
     '1. Không tranh luận cảm tính — "em thấy khách hàng rất quan tâm" không bằng số.\n'
     '2. Nhìn điểm nghẽn — bước nào có tỷ lệ rơi cao nhất? Đó là ưu tiên sửa.\n'
     '3. Chọn thử nghiệm tiếp theo — dựa trên dữ liệu, chọn giả thuyết để test.\n\n'
     '[TƯƠNG TÁC] Ví dụ: tiếp cận nhiều nhưng mua ít. '
     'Câu hỏi đặt ra: vấn đề nằm ở thông điệp, giá hay trải nghiệm dùng thử?'),
]

for num, title, section_name, seconds, content in slides_tt:
    slide_header(doc, str(num).zfill(2), title, section_name)
    note(doc, f"⏱  Khoảng {seconds} giây")
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("[TƯƠNG TÁC]"):
            action(doc, line, ACCENT)
        elif line.startswith("[CHUYỂN CẢNH]"):
            action(doc, line, TEAL)
        elif line.startswith('"'):
            talk(doc, line)
        else:
            body(doc, line)

# ─── PHẦN 5 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 5  ·  TÀI CHÍNH  (Slide 29–37)  ·  ~4 phút")
add_hr(doc)

slides_tc = [
    (29, "Chi phí khởi tạo", "Finance", 23,
     '"Chi phí khởi tạo phải được chia theo nhóm quyết định, '
     'không gom thành một cục mơ hồ."\n\n'
     'Năm nhóm chi phí:\n'
     '1. Một lần — pháp lý, thiết bị. Chi một lần ngay từ đầu.\n'
     '2. Hàng tháng — mặt bằng, nhân sự, vận hành. Chi định kỳ.\n'
     '3. Dự phòng — luôn để ra 20% cho phát sinh.\n'
     '4. Có thể trì hoãn — marketing, mở rộng. Chưa cần chi ngay.\n\n'
     '"Biết khoản nào bắt buộc, khoản nào chờ được. '
     'Không chi trì hoãn trước bắt buộc."'),
    (30, "Dòng tiền ban đầu", "Finance", 25,
     '"Lợi nhuận trên giấy không cứu được doanh nghiệp hết tiền mặt."\n\n'
     'Bốn thành phần cần theo dõi:\n'
     '1. Tiền vào — doanh thu thực nhận, không phải đơn hàng.\n'
     '2. Tiền ra — chi phí thực chi, không phải chi phí dự kiến.\n'
     '3. Runway — số tháng còn sống nếu không có thêm tiền vào.\n'
     '4. Khoản phải thu — đã bán nhưng chưa thu tiền. Đây là tiền trên giấy.\n\n'
     '[TƯƠNG TÁC] Hỏi lớp: có doanh thu 50 triệu nhưng 40 triệu là nợ phải thu — '
     'doanh nghiệp có khá hơn không?'),
    (31, "Điểm hòa vốn", "Finance", 25,
     '"Hòa vốn trả lời câu hỏi: cần bán bao nhiêu để không mất thêm tiền."\n\n'
     'Bốn yếu tố xác định điểm hòa vốn:\n'
     '1. Chi phí cố định — tiền thuê, lương, không đổi hàng tháng.\n'
     '2. Biến phí — chi phí tăng theo sản lượng, nguyên vật liệu.\n'
     '3. Giá bán — giá mỗi đơn vị sản phẩm.\n'
     '4. Sản lượng hòa vốn — số sản phẩm cần bán để doanh thu = chi phí.\n\n'
     '"Không cần đi sâu công thức. Hòa vốn là mốc giúp founder biết mục tiêu bán hàng tối thiểu."'),
    (32, "Bảng phân bổ nguồn lực", "Finance", 22,
     '"Nguồn lực ít nên phải phân bổ theo mục tiêu học và sống sót."\n\n'
     'Nhìn vào biểu đồ tròn:\n'
     'Sản phẩm chiếm phần lớn nhất — vì đây là nền tảng.\n'
     'Bán hàng chiếm vị trí thứ hai — vì không bán được thì không có tiền.\n'
     'Vận hành, Tài chính, Pháp lý — chiếm phần còn lại.\n\n'
     '[TƯƠNG TÁC] Trade-off: không thể vừa làm sản phẩm hoàn hảo, '
     'vừa marketing lớn, vừa thuê nhiều người, vừa giữ runway dài — '
     'nếu nguồn lực nhỏ.'),
    (33, "Rủi ro tài chính", "Finance", 22,
     '"Rủi ro tài chính nguy hiểm vì nó âm thầm tích tụ trước khi lộ ra."\n\n'
     'Bốn rủi ro thường gặp:\n'
     '1. Định giá sai — bán giá thấp hơn chi phí thật, lỗ âm thầm.\n'
     '2. Burn rate cao — chi tiêu hàng tháng cao hơn dự kiến, runway ngắn hơn.\n'
     '3. Thu tiền chậm — khách nợ kéo dài, dòng tiền bị thắt.\n'
     '4. Không có dự phòng — chi phí phát sinh không có tiền trả.\n\n'
     '"Bán nhiều nhưng biên lợi nhuận thấp, khách nợ tiền, tồn kho cao. '
     'Nhìn doanh thu thôi là chưa đủ."'),
    (34, "Quản trị rủi ro", "Finance", 22,
     '"Quản trị rủi ro là biến nỗi lo thành danh sách hành động."\n\n'
     'Năm bước:\n'
     '1. Nhận diện — liệt kê tất cả rủi ro có thể xảy ra.\n'
     '2. Ước lượng — xác suất bao nhiêu %, tác động lớn hay nhỏ?\n'
     '3. Ưu tiên — tập trung vào rủi ro cao nhất trước.\n'
     '4. Phản ứng — tránh, giảm, chuyển hay chấp nhận?\n'
     '5. Theo dõi — rủi ro thay đổi theo thời gian, cần cập nhật.\n\n'
     '"Không cần quá học thuật. Quan trọng là nhóm biết '
     'rủi ro nào tránh, rủi ro nào giảm, rủi ro nào chấp nhận."'),
    (35, "Rủi ro thị trường", "Finance", 22,
     '"Thị trường không sai. Giả định của founder mới dễ sai."\n\n'
     'Bốn rủi ro thị trường:\n'
     '1. Nhu cầu yếu — giả định khách cần nhưng thực tế không đủ nhu cầu.\n'
     '2. Giá không phù hợp — định giá cao hơn giá trị khách nhận được.\n'
     '3. Kênh bán đắt — chi phí tiếp cận khách cao hơn giá trị họ mang lại.\n'
     '4. Đối thủ phản ứng — đối thủ giảm giá hoặc copy sau khi mình launch.\n\n'
     '"Khách không mua không có nghĩa khách chưa hiểu mình. '
     'Có thể giả định giá trị của mình chưa đúng."'),
    (36, "Rủi ro vận hành", "Finance", 22,
     '"Vận hành yếu làm trải nghiệm khách không ổn định — '
     'dù sản phẩm có nhu cầu."\n\n'
     'Bốn rủi ro:\n'
     '1. Giao trễ — không giao đúng hẹn, khách mất tin.\n'
     '2. Sai chất lượng — sản phẩm không đúng như cam kết.\n'
     '3. Thiếu owner — không ai chịu trách nhiệm khi có vấn đề.\n'
     '4. Không ghi nhận lỗi — cùng một lỗi lặp lại nhiều lần vì không có hệ thống theo dõi.\n\n'
     '"Vận hành là nơi lời hứa gặp thực tế. SOP giúp giảm rủi ro này."'),
    (37, "Ma trận rủi ro", "Finance", 23,
     '"Ma trận giúp chọn rủi ro cần xử lý trước."\n\n'
     'Từng điểm rủi ro trên ma trận:\n'
     'Đốt tiền nhanh — tác động cao, xác suất cao → ưu tiên số một.\n'
     'Không có khách — tác động cao, xác suất cao → ưu tiên số hai.\n'
     'Sai loại hình pháp lý — tác động cao, xác suất trung bình.\n'
     'Không đọc số liệu — tác động cao, xác suất cao.\n'
     'Founder ôm việc — tác động trung bình, xác suất rất cao.\n\n'
     '"Góc phần tư trên bên phải = ưu tiên đầu tiên. '
     'Không phải rủi ro nào cũng xử lý như nhau."'),
]

for num, title, section_name, seconds, content in slides_tc:
    slide_header(doc, str(num).zfill(2), title, section_name)
    note(doc, f"⏱  Khoảng {seconds} giây")
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("[TƯƠNG TÁC]"):
            action(doc, line, ACCENT)
        elif line.startswith("[CHUYỂN CẢNH]"):
            action(doc, line, TEAL)
        elif line.startswith('"'):
            talk(doc, line)
        else:
            body(doc, line)

# ─── PHẦN 6 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 6  ·  KIỂM SOÁT & KPI  (Slide 38–39)  ·  ~1 phút")
add_hr(doc)

slides_ks = [
    (38, "KPI giai đoạn khởi sự", "Control", 22,
     '"KPI giai đoạn đầu phải đo hai thứ: học nhanh và sống sót — '
     'không phải tăng trưởng."\n\n'
     'Năm KPI cốt lõi:\n'
     '1. Doanh thu — đo bằng tiền thật, không phải đơn hàng.\n'
     '2. Biên gộp — lợi nhuận trên mỗi đơn hàng.\n'
     '3. Chuyển đổi — tỷ lệ khách tiếp cận thành khách mua.\n'
     '4. Runway — số tuần còn sống.\n'
     '5. Mua lại — khách quay lại mua lần hai, ba.\n\n'
     '"Chọn số khiến founder ra quyết định hằng tuần: '
     'bán gì, cắt gì, thử gì tiếp."'),
    (39, "Dashboard KPI mẫu", "Control", 25,
     '"Dashboard tốt là bảng điều khiển, không phải nơi trang trí số liệu."\n\n'
     'Bốn KPI bắt buộc với ngưỡng cảnh báo:\n'
     '1. Doanh thu tuần >= 18 triệu — Tín hiệu: tăng hay giảm theo kênh nào?\n'
     '2. Biên gộp >= 38% — Tín hiệu: giá vốn có lệch không?\n'
     '3. Chuyển đổi >= 6.2% — Tín hiệu: thông điệp bán hàng có rõ không?\n'
     '4. Cash runway >= 10 tuần — Tín hiệu: có cần cắt chi phí không?\n\n'
     '"Mỗi KPI có owner. Owner biết khi nào cần hành động."'),
]

for num, title, section_name, seconds, content in slides_ks:
    slide_header(doc, str(num).zfill(2), title, section_name)
    note(doc, f"⏱  Khoảng {seconds} giây")
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("[TƯƠNG TÁC]"):
            action(doc, line, ACCENT)
        elif line.startswith("[CHUYỂN CẢNH]"):
            action(doc, line, TEAL)
        elif line.startswith('"'):
            talk(doc, line)
        else:
            body(doc, line)

# ─── PHẦN 7 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 7  ·  THỰC THI 90 NGÀY  (Slide 40–43)  ·  ~1.5 phút")
add_hr(doc)

slides_tt90 = [
    (40, "Timeline 90 ngày đầu", "Execution", 22,
     '"90 ngày đầu nên được chia thành ba nhịp rõ ràng:"\n\n'
     '30 ngày — Setup: Dựng nền. Pháp lý xong, vai trò rõ, MVP chạy thử.\n'
     '60 ngày — Validate: Kiểm chứng. Có khách thật, có dữ liệu phễu, quy trình được chỉnh.\n'
     '90 ngày — Stabilize: Ổn định. Runway còn đủ, KPI bắt đầu đạt, quyết định tiếp theo rõ.\n\n'
     '"Không phải mọi doanh nghiệp giống nhau. Nhưng 30-60-90 giúp kiểm soát kỳ vọng."'),
    (41, "Mốc 30 ngày", "Execution", 19,
     '30 ngày đầu: khóa nền tảng. Bốn mục tiêu:\n\n'
     '1. Pháp lý xong — hồ sơ đăng ký, thuế, hóa đơn, ngân hàng.\n'
     '2. Owner rõ — mỗi người biết việc của mình, biết deadline của mình.\n'
     '3. MVP chạy — có sản phẩm đủ để khách thật dùng và trả tiền.\n'
     '4. Cash sheet có số — biết đang tiêu bao nhiêu, còn bao lâu.\n\n'
     '"Không cần scale. Cần chứng minh doanh nghiệp có thể chạy hợp lệ."'),
    (42, "Mốc 60 ngày", "Execution", 19,
     '60 ngày: giai đoạn kiểm chứng. Bốn mục tiêu:\n\n'
     '1. Có khách thật — không phải bạn bè, không phải thử nghiệm. Khách trả tiền thật.\n'
     '2. Có dữ liệu phễu — biết khách rơi ở bước nào.\n'
     '3. SOP được sửa — quy trình đầu tiên đã có phiên bản 1.0, giờ tinh chỉnh.\n'
     '4. Kênh bán rõ hơn — biết kênh nào hiệu quả, kênh nào cần bỏ.\n\n'
     '"Nếu sau 60 ngày vẫn chỉ nói cảm giác, nhóm đang thiếu hệ thống đo."'),
    (43, "Mốc 90 ngày", "Execution", 19,
     '90 ngày: lúc quyết định lớn. Bốn câu hỏi cần trả lời:\n\n'
     '1. Runway còn bao lâu? — Nếu dưới 8 tuần, cần hành động ngay.\n'
     '2. KPI nào đạt? — Nếu biên gộp và chuyển đổi không cải thiện, cần xem lại.\n'
     '3. Giả định nào sai? — Mỗi sai lệch là một bài học, không phải thất bại.\n'
     '4. Quyết định tiếp theo là gì? — Tiếp tục, chỉnh hướng, hay thu hẹp?\n\n'
     '"Biết dừng hoặc pivot cũng là năng lực triển khai. '
     'Không phải cứ làm tiếp là tốt."'),
]

for num, title, section_name, seconds, content in slides_tt90:
    slide_header(doc, str(num).zfill(2), title, section_name)
    note(doc, f"⏱  Khoảng {seconds} giây")
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("[TƯƠNG TÁC]"):
            action(doc, line, ACCENT)
        elif line.startswith("[CHUYỂN CẢNH]"):
            action(doc, line, TEAL)
        elif line.startswith('"'):
            talk(doc, line)
        else:
            body(doc, line)

# ─── PHẦN 8 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 8  ·  RECAP & TƯƠNG TÁC  (Slide 44–48)  ·  ~3 phút")
add_hr(doc)

# Slide 44
slide_header(doc, "44", "Những sai lầm thường gặp", "Recap")
note(doc, "⏱  Khoảng 23 giây")
body(doc, '"Sai lầm lớn nhất: triển khai bằng cảm hứng thay vì bằng hệ thống."')
body(doc, "")
talk(doc, '"1. Chọn pháp lý theo cảm tính — chọn cái tiện lúc đầu, vướng sau.  → Slide 8, 9."')
talk(doc, '"2. Không phân vai — founder ôm hết, kiệt sức.  → Slide 15, 16."')
talk(doc, '"3. Không đo dòng tiền — có doanh thu nhưng hết tiền mặt.  → Slide 30."')
talk(doc, '"4. Làm marketing trước khi có giả thuyết — chi tiền mà không biết đo cách nào.  → Slide 23, 26."')
note(doc, "Mỗi lỗi gắn với slide đã trình bày trước đó. Câu chuyện liên kết.")
action(doc, "[CHUYỂN CẢNH] →  \"Trước khi vận hành thật, cần checklist cuối cùng.\"")

# Slide 45
slide_header(doc, "45", "Checklist trước khi vận hành", "Recap")
note(doc, "⏱  Khoảng 27 giây")
body(doc, '"Trước khi chạy thật, hãy kiểm tra năm câu hỏi sống còn. '
       'Slide này có thể chụp lại."')
body(doc, "")
talk(doc, '"Câu 1: Hợp lệ chưa? — Pháp lý đã xong chưa?"')
talk(doc, '"Câu 2: Ai làm gì? — Vai trò đã phân rõ chưa?"')
talk(doc, '"Câu 3: Bán cho ai? — Khách hàng mục tiêu đã xác định chưa?"')
talk(doc, '"Câu 4: Tiền còn bao lâu? — Runway còn đủ không?"')
talk(doc, '"Câu 5: Đo bằng gì? — KPI đã chọn và theo dõi chưa?"')
action(doc, "[TƯƠNG TÁC] →  \"Thầy ơi, em muốn dành ít phút để lớp trả lời nhanh 5 câu hỏi này cho một tình huống giả định.\"")

# Slide 46
slide_header(doc, "46", "Tình huống tương tác với lớp", "Interaction")
note(doc, "⏱  Khoảng 35 giây  ·  ĐÂY LÀ SLIDE TƯƠNG TÁC CHÍNH")
body(doc, "")
body(doc, "Dành 30-35 giây cho phần này. Không vội.")
body(doc, "")
body(doc, "Tình huống:", bold=True)
talk(doc, '"Một nhóm muốn mở thương hiệu đồ ăn healthy quanh trường. '
       'Câu hỏi: Bắt đầu từ đâu?"')
body(doc, "")
body(doc, "[TƯƠNG TÁC] Dừng 15 giây, cho lớp suy nghĩ. Có thể giơ tay hoặc thảo luận nhóm.")
body(doc, "")
body(doc, "Gợi ý:")
action(doc, "Dùng framework của chương — pháp lý, MVP, kênh bán, dòng tiền, KPI. "
          "Ưu tiên cái nào trước? Không cần đáp án duy nhất. "
          "Mục tiêu là buộc người nghe dùng framework thay vì đoán.")

# Slide 47
slide_header(doc, "47", "Câu hỏi thực tế 1", "Interaction")
note(doc, "⏱  Khoảng 27 giây")
body(doc, "Tình huống:", bold=True)
talk(doc, '"Nhóm có 20 triệu và muốn bán đồ ăn healthy quanh trường trong 30 ngày. '
       'Việc đầu tiên nên làm là gì?"')
body(doc, "")
talk(doc, "A. Làm logo và fanpage")
talk(doc, "B. Test 10 khách đầu tiên")
talk(doc, "C. Thuê mặt bằng nhỏ")
talk(doc, "D. Đăng ký pháp lý đầy đủ ngay")
body(doc, "")
body(doc, "[TƯƠNG TÁC] Cho lớp chọn A, B, C hoặc D.", bold=True)
body(doc, "")
body(doc, "Sau đó gợi ý:", muted=True)
action(doc, "Với vốn nhỏ và 30 ngày, ưu tiên thường là B — test khách thật — "
          "kết hợp hiểu rủi ro pháp lý tối thiểu. "
          "Đăng ký đầy đủ tốn thời gian và chi phí, trong khi chưa biết khách có mua không.")

# Slide 48
slide_header(doc, "48", "Câu hỏi thực tế 2", "Interaction")
note(doc, "⏱  Khoảng 27 giây")
body(doc, "Tình huống:", bold=True)
talk(doc, '"Sau 2 tuần launch: 100 người hỏi nhưng chỉ 6 người mua. '
       'Bạn sửa phần nào trước?"')
body(doc, "")
talk(doc, "A. Giảm giá ngay")
talk(doc, "B. Đổi thông điệp bán hàng")
talk(doc, "C. Đổi sản phẩm")
talk(doc, "D. Đổi kênh bán")
body(doc, "")
body(doc, "[TƯƠNG TÁC] Cho lớp chọn một hướng.", bold=True)
body(doc, "")
body(doc, "Sau đó gợi ý:", muted=True)
action(doc, "100 hỏi nhưng 6 mua = rơi ở bước quan tâm → thông điệp bán hàng. "
          "Nhưng quan trọng là — phải xem DỮ LIỆU trước khi quyết định. "
          "Hỏi: khách hỏi gì? Họ so sánh với ai? Họ phản đối ở đâu?")

# ─── PHẦN 9 ────────────────────────────────────────────────────────────────
h1(doc, "PHẦN 9  ·  KẾT  (Slide 49–50)  ·  ~1 phút")
add_hr(doc)

# Slide 49
slide_header(doc, "49", "Recap toàn chương", "Closing")
note(doc, "⏱  Khoảng 25 giây")
body(doc, '"Recap toàn chương bằng bốn ý chính. Không thêm kiến thức mới. Đọc chậm, có ngữ điệu."')
body(doc, "")
talk(doc, '"1. Pháp lý để hợp lệ — chọn loại hình đúng, đăng ký đủ, tránh rủi ro."')
talk(doc, '"2. Vận hành để lặp lại — phân vai, quy trình, SOP, không phụ thuộc trí nhớ."')
talk(doc, '"3. Tài chính để sống sót — theo dõi dòng tiền, biết runway, tính hòa vốn."')
talk(doc, '"4. KPI để học nhanh — đo hành vi, quyết định bằng dữ liệu, không bằng cảm tính."')
action(doc, "[CHUYỂN CẢNH] →  \"Và kết thúc bằng một câu punchline.\"")

# Slide 50
slide_header(doc, "50", "Closing punchline", "Closing")
note(doc, "⏱  Khoảng 19 giây  ·  ĐÂY LÀ SLIDE CUỐI")
body(doc, "")
body(doc, "DỪNG MỘT NHỊP TRƯỚC KHI NÓI. Nhìn lớp. Im lặng 2 giây.", bold=True)
body(doc, "")
talk(doc, '"Một doanh nghiệp không bắt đầu khi có ý tưởng."')
body(doc, "")
body(doc, "DỪNG 1 GIÂY", muted=True)
body(doc, "")
talk(doc, '"Mà khi hệ thống đầu tiên chạy được."')
body(doc, "")
body(doc, "DỪNG 1 GIÂY", muted=True)
body(doc, "")
talk(doc, '"Hợp lệ. Chạy được. Đo được."')
body(doc, "")
body(doc, "DỪNG 1 GIÂY", muted=True)
body(doc, "")
talk(doc, '"Cảm ơn thầy và các bạn đã lắng nghe. Nhóm 3 sẵn sàng Q&A."')
note(doc, "Tip hướng nội: câu punchline là điểm nhấn cuối. Nói CHẬM, rõ, có ngừng. "
     "Không cần hét hay làm drama. Cứ bình tĩnh, tự tin — đó mới là phong cách đáng nhớ.")
note(doc, "Bấm Q để mở panel Q&A dự phòng trên màn hình.")

# ─── BẢNG TỔNG HỢP ────────────────────────────────────────────────────────
add_hr(doc)
h1(doc, "BẢNG TỔNG HỢP THỜI GIAN")
add_hr(doc)

summary = [
    ("1", "Cover", "16s"),
    ("2", "Thành viên nhóm", "16s"),
    ("3", "Vì sao Chương 6 quan trọng", "20s"),
    ("4", "Bản đồ toàn chương", "20s"),
    ("5", "Từ ý tưởng sang DN thật", "17s"),
    ("6", "Checklist triển khai", "19s"),
    ("7", "Chọn mô hình kinh doanh", "19s"),
    ("8", "Chọn loại hình pháp lý", "22s"),
    ("9", "Decision tree", "27s"),
    ("10", "Hồ sơ đăng ký", "19s"),
    ("11", "Quy trình đăng ký", "23s"),
    ("12", "Chi phí pháp lý", "19s"),
    ("13", "Rủi ro pháp lý", "22s"),
    ("14", "Case chọn sai loại hình", "27s"),
    ("15", "Bộ máy vận hành", "22s"),
    ("16", "Vai trò founder", "20s"),
    ("17", "Vai trò nhân sự lõi", "20s"),
    ("18", "Sơ đồ tổ chức", "22s"),
    ("19", "Quy trình vận hành tối thiểu", "23s"),
    ("20", "SOP", "23s"),
    ("21", "Sản phẩm tối thiểu", "22s"),
    ("22", "Lean Startup và MVP", "23s"),
    ("23", "Build Measure Learn", "25s"),
    ("24", "Kiểm chứng khách hàng", "23s"),
    ("25", "Kênh bán hàng", "22s"),
    ("26", "Marketing launch plan", "23s"),
    ("27", "Phễu khách hàng", "22s"),
    ("28", "Biểu đồ chuyển đổi", "22s"),
    ("29", "Chi phí khởi tạo", "23s"),
    ("30", "Dòng tiền ban đầu", "25s"),
    ("31", "Điểm hòa vốn", "25s"),
    ("32", "Phân bổ nguồn lực", "22s"),
    ("33", "Rủi ro tài chính", "22s"),
    ("34", "Quản trị rủi ro", "22s"),
    ("35", "Rủi ro thị trường", "22s"),
    ("36", "Rủi ro vận hành", "22s"),
    ("37", "Ma trận rủi ro", "23s"),
    ("38", "KPI giai đoạn khởi sự", "22s"),
    ("39", "Dashboard KPI", "25s"),
    ("40", "Timeline 90 ngày", "22s"),
    ("41", "Mốc 30 ngày", "19s"),
    ("42", "Mốc 60 ngày", "19s"),
    ("43", "Mốc 90 ngày", "19s"),
    ("44", "Sai lầm thường gặp", "23s"),
    ("45", "Checklist trước vận hành", "27s"),
    ("46", "Tình huống tương tác", "35s"),
    ("47", "Câu hỏi thực tế 1", "27s"),
    ("48", "Câu hỏi thực tế 2", "27s"),
    ("49", "Recap toàn chương", "25s"),
    ("50", "Closing punchline", "19s"),
    ("", "TỔNG CỘNG", "~18-20 phút"),
]

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("Trang")
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = DARK
run2 = p.add_run("            " + "Tên slide")
run2.font.size = Pt(9)
run2.font.bold = True
run2.font.color.rgb = DARK
run3 = p.add_run("                              " + "Thời gian")
run3.font.size = Pt(9)
run3.font.bold = True
run3.font.color.rgb = DARK

for num, title, time_str in summary:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if num == "":
        run = p.add_run(f"{title}")
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = ACCENT
    else:
        run = p.add_run(f"{num.zfill(2)}")
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        run2 = p.add_run(f"    {title}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = DARK
        run3 = p.add_run(f"                        {time_str}")
        run3.font.size = Pt(9)
        run3.font.color.rgb = GRAY

# ─── PHÍM TẮT ────────────────────────────────────────────────────────────────
add_hr(doc)
h1(doc, "PHÍM TẮT KHI THUYẾT TRÌNH")
add_hr(doc)

shortcuts = [
    ("→  hoặc  Space", "Slide tiếp theo"),
    ("←", "Slide trước"),
    ("F", "Toàn màn hình"),
    ("S", "Mở speaker notes (cửa sổ riêng cho người thuyết trình)"),
    ("O", "Overview tất cả slide"),
    ("Q", "Mở panel Q&A dự phòng"),
    ("ESC", "Thoát toàn màn hình"),
]

for key, desc in shortcuts:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run1 = p.add_run(f"  {key}")
    run1.font.size = Pt(10)
    run1.font.bold = True
    run1.font.color.rgb = ACCENT
    run2 = p.add_run(f"   {desc}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK

# ─── MẸO THUYẾT TRÌNH ────────────────────────────────────────────────────
add_hr(doc)
h1(doc, "MẸO THUYẾT TRÌNH CHO NGƯỜI HƯỚNG NỘI")
add_hr(doc)

tips = [
    ("Trước khi nói",
     "Đứng vững, nhìn thẳng. Im lặng 1-2 giây trước khi bắt đầu là cách tạo ấn tượng tốt "
     "mà không cần nói gì. Không cần phải vội bắt đầu."),
    ("Giọng nói",
     "Nói chậm ở câu quan trọng, nói đều ở câu phụ. Có ngừng nhịp sau mỗi ý. "
     "Không cần phải lớn tiếng — giọng bình tĩnh, rõ ràng còn ấn tượng hơn giọng hét."),
    ("Khi lúng túng",
     "Nhìn slide, đọc bullet → tự tin hơn. Im lặng 2-3 giây vẫn OK, không ai phán xét. "
     "Nếu quên: 'Em xin phép nhắc lại...' rồi tiếp tục. Không cần phải hoàn hảo."),
    ("Tương tác",
     "Không cần ai trả lời hoàn hảo. Tự trả lời sau 5-7 giây là bình thường. "
     "Khen ngắn nếu có câu trả lời tốt: 'Cảm ơn bạn, đúng rồi!'"),
    ("Chuyển cảnh",
     "Luôn nói một câu chuyển trước khi bấm slide mới. "
     "Câu chuyển gợi tò mò cho slide tiếp theo, ví dụ: 'Vậy bước tiếp theo là gì?'"),
]

for title_tip, content_tip in tips:
    h3(doc, title_tip)
    body(doc, content_tip)

# ─── SAVE ───────────────────────────────────────────────────────────────────
output_path = "d:/Study/Nam 3 ky 3/Khởi nghiệp/WebPresentation/web-ppt/SCRIPT_THUYET_TRINH.docx"
doc.save(output_path)
print("OK - Saved:", output_path)
